import os
from datetime import datetime
import streamlit as st
from io import BytesIO
from PIL import Image
from docx import Document
from docx.shared import Inches
import google.generativeai as genai
from dotenv import load_dotenv
from tensorflow.keras.models import load_model
import numpy as np
import streamlit as st
api_key = st.secrets["GEMINI_API_KEY"]
# Charger le modèle CNN entraîné
cnn_model = load_model("CNN_model1.h5")

# Classes (à adapter selon ton dataset)
class_labels = ['Abnormal', 'Covid-19', 'Normal', 'MI', 'HMI', 'AHB']

# Charger les variables d'environnement
api_key = st.secrets["GEMINI_API_KEY"]
if api_key is None:
    raise ValueError("GEMINI_API_KEY is not set in environment variables")
genai.configure(api_key=api_key)

# Paramètres du modèle
generation_config = {
    "temperature": 1,
    "top_p": 0.95,
    "top_k": 64,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain",
}

model = genai.GenerativeModel(
    model_name="gemini-1.5-flash",
    generation_config=generation_config,
)
def predict_ecg_class(image_file):
    image = Image.open(image_file).convert('RGB').resize((224, 224))  
    image_array = np.array(image) / 255.0
    image_array = np.expand_dims(image_array, axis=0)  
    prediction = cnn_model.predict(image_array)
    predicted_class = class_labels[np.argmax(prediction)]
    return predicted_class

# Fonction pour analyser une image ECG et générer un rapport
def generate_ecg_details(ecg_image):
    image = Image.open(ecg_image)
    current_date = datetime.now().strftime('%Y-%m-%d')

    prompt = f"""Analyze this ECG image and provide a detailed report. Follow this structure:

**ECG ANALYSIS REPORT**

**1. PATIENT INFORMATION:**
- Name:
- Age:
- Gender:
- ID Number:
- Date of ECG:

**2. CLINICAL INFORMATION:**
- Reason for ECG:
- Relevant Medical History:
- Medications:

**3. ECG TECHNICAL DETAILS:**
- ECG Machine Used:
- Lead Configuration:
- Calibration:
- Recording Quality:

**4. ECG FINDINGS:**
**Rhythm and Rate:**
- Heart Rate:
- Rhythm:
- P Waves:
- PR Interval:
- QRS Complex:
- QT/QTc Interval:
- ST Segment:
- T Waves:

**Axis:**
- P Wave Axis:
- QRS Axis:
- T Wave Axis:

**Conduction and Morphology:**
- Atrial Conduction:
- Ventricular Conduction:
- QRS Morphology:
- ST-T Changes:

**5. INTERPRETATION:**
- Normal or Abnormal:
- Diagnosis/Findings:
- Comparison with Previous ECG (if available):

**6. CONCLUSION AND RECOMMENDATIONS:**
- Summary:
- Recommendations:

**7. REPORTING CARDIOLOGIST:**
- Name:
- Signature: Unable to provide signature for AI-generated report.
- Date of Report: {current_date}
"""

    chat_session = model.start_chat(history=[])
    predicted_class = predict_ecg_class(ecg_image)
    full_prompt = prompt + f"\n\n**Anomaly Class (Predicted by Our Model):** {predicted_class}\n\nNow complete the rest of the report using the above prediction as reference."
    response = chat_session.send_message([full_prompt, image])
    return response.text

# Fonction pour créer un document Word contenant le rapport ECG
def create_doc(report_text, ecg_image):
    doc = Document()
    doc.add_heading('ECG ANALYSIS REPORT', 0)

    for line in report_text.split("\n"):
        if line.strip() == '':
            continue
        if line.startswith('**') and line.endswith('**'):
            doc.add_heading(line.strip('**'), level=1)
        elif line.startswith('-'):
            doc.add_paragraph(line.strip(), style='List Bullet')
        else:
            doc.add_paragraph(line.strip())

    doc.add_heading('ECG Tracing:', level=1)
    image_stream = BytesIO(ecg_image.getvalue())
    doc.add_picture(image_stream, width=Inches(6))

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# Interface utilisateur avec Streamlit
def main():
    st.title("🫀Heart Health Chatbot - Get Instant ECG Analysis")

    # Section Upload ECG
    st.header("Upload ECG Image")
    ecg_image = st.file_uploader("Upload an ECG Image", type=["png", "jpg", "jpeg"])

    if ecg_image is not None:
        st.image(ecg_image, caption='Uploaded ECG Image', use_column_width=True)

        if st.button("Generate ECG Report"):
            with st.spinner("Analyzing ECG image..."):
                ecg_details = generate_ecg_details(ecg_image)
            st.header("Generated ECG Report")
            st.markdown(ecg_details)

            # Stocker le rapport dans la session pour le téléchargement
            st.session_state.ecg_details = ecg_details

        # Bouton de téléchargement du rapport
        if hasattr(st.session_state, 'ecg_details'):
            doc_file_stream = create_doc(st.session_state.ecg_details, ecg_image)
            st.download_button(
                label="Download ECG Report",
                data=doc_file_stream,
                file_name="ECG_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    # Section Chatbot IA
    st.header("Ask Your AI Cardiologist ")

    if "messages" not in st.session_state:
        st.session_state.messages = []

    # Afficher l'historique du chat
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # Saisie utilisateur
    user_input = st.chat_input("Ask me anything...")
    if user_input:
        st.session_state.messages.append({"role": "user", "content": user_input})
        
        with st.chat_message("user"):
            st.markdown(user_input)

        with st.spinner("Thinking..."):
            chat_session = model.start_chat(history=[])
            response = chat_session.send_message(user_input)
            bot_response = response.text

        st.session_state.messages.append({"role": "assistant", "content": bot_response})

        with st.chat_message("assistant"):
            st.markdown(bot_response)

if __name__ == '__main__':
    main()
